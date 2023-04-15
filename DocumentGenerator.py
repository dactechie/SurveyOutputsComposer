from typing import Generator, List, Any
from docxcompose.composer import Composer
from docx import Document
from docx.text.paragraph import Paragraph
from SharepointHandler import SharepointHandler

from docx_fetch.GenericDocx import GenericDocument
from docx_fetch.LocalDocx import LocalDocx
from docx_fetch.SharepointDocx import SharepointDocx
from utils import secret
from utils.text_util import get_variables_in_text
import constants as cns

current_version_all_docs = 1.0 # temporary. Each doc would have its current version

root_template_folder = 'template_fragments' #TODO: get rid of this from here
RUNTIME_STAGE ="Staging" # Prod
SP_PATH_PREFIX = f'/Shared Documents/ICT/ATOM/{RUNTIME_STAGE}/{root_template_folder}' # TODO move this

def get_paras(doc:Document) -> (Generator[Paragraph, None, None], List[Any]):# type: ignore
  for table in doc.tables:
    for row in table.rows:
      for cell in row.cells:
        for para in cell.paragraphs:
          variables = get_variables_in_text(para.text)
          if variables:
            yield para, variables
  
def replace_placeholders(normalised_survey_data: dict, doc: Document) -> Generator[Document, None, None]:  # type: ignore
  for para, para_variables in get_paras(doc):
    for pv in para_variables:
      sdata = normalised_survey_data.get(pv)
      if not sdata:
        print(f"No Data for {pv}.")
        continue
      para.text = para.text.replace(f'%%{pv}%%', sdata)
  return doc

# TODO: handle multiple DOC versions
def get_sp_docs(sph:SharepointHandler, path_versions, sp_prefix)-> List[SharepointDocx]:
  sp_docs: List[SharepointDocx] = [] 
  for fp, version in path_versions:    
    if version != current_version_all_docs:
      full_path = f"{sp_prefix}/{fp}"
      sp_docs.append(SharepointDocx(full_path, sph))
    
  return sp_docs
    

def build_doc_list(ordered_fragments_paths, versions) -> List[GenericDocument]:  
  documents : list[GenericDocument]= []
  sp_docs: list[SharepointDocx] = []
  path_versions = zip(ordered_fragments_paths, versions)
  if any(v for v in versions if v != current_version_all_docs):    
    sph = SharepointHandler(cns.site_url,
                            client_id=secret.get_value_for_key(cns.KEY_SHAREPOINT_CLIENTID),
                            client_secret=secret.get_value_for_key(cns.KEY_SHAREPOINT_SECRET))
     
    sp_docs = list(reversed(get_sp_docs(sph, path_versions, SP_PATH_PREFIX)))

  for fp, version in zip(ordered_fragments_paths, versions):     
    docx_fetcher = None
    if version == current_version_all_docs:
      full_path = f"./{root_template_folder}/{fp}"
      docx_fetcher = LocalDocx(full_path)
    else:
      docx_fetcher = sp_docs.pop()

    doc = docx_fetcher.get_document()
    documents.append(doc)
  return documents


class DocumentGenerator:

  def __init__(self, base_doc_path):
    # self.base_doc = Document(base_doc_path)
    self.composer = Composer(Document(base_doc_path))


  def _compose_document(self, ordered_fragments: List[GenericDocument], survey_data: dict) -> Composer:
 
    # survey_data = get_normalised_dict(survey_data)
    for fragment in ordered_fragments:
      #replace variable placeholders here
          
      replaced_fragment = replace_placeholders(survey_data, fragment)
      self.composer.append(replaced_fragment)

    return self.composer


  def build_document(self, documents, 
          survey_data) -> Composer:
    """
      A. Composes Document:
        Joins all fragments after doing the following:
        1. Normalises SurveyData Dictionary : Flattens Dict by concat'ing parent+child keys
            (which matche the doc template variables)
        2. Does variable replacement from the survey data

      B. Constructs the output file name
    """             
    final_doc = self._compose_document(documents, survey_data)
   
    return final_doc
